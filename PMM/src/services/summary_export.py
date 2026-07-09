"""Summary export helpers for TXT, DOCX, and PDF downloads."""

import html
import json
import logging
import os
import re
from io import BytesIO


logger = logging.getLogger(__name__)

SECTION_ALIASES = {
    'overview': ['overview', 'summary', 'meeting_summary', 'executive_summary'],
    'key_points': ['key_points', 'key_discussion_points', 'discussion_points', 'key_discussions', 'topics'],
    'action_items': ['action_items', 'actions', 'tasks', 'todos', 'to_dos'],
    'decisions': ['decisions', 'decision_points'],
    'next_steps': ['next_steps', 'nextsteps', 'follow_ups', 'followups'],
    'risks': ['risks', 'blockers', 'concerns'],
}


def _clean_text(value):
    if value is None:
        return ''
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float, bool)):
        return str(value)
    if isinstance(value, dict):
        return ' - '.join(
            f'{_humanize_key(key)}: {_clean_text(item)}'
            for key, item in value.items()
            if _clean_text(item)
        )
    return str(value).strip()


def _humanize_key(key):
    return re.sub(r'\s+', ' ', str(key).replace('_', ' ').replace('-', ' ')).strip().title()


def _strip_code_fence(text):
    text = (text or '').strip()
    match = re.match(r'^```(?:json)?\s*([\s\S]*?)\s*```$', text, flags=re.IGNORECASE)
    return match.group(1).strip() if match else text


def _normalize_json_like_text(text):
    return (
        str(text or '')
        .replace('\u201c', '"')
        .replace('\u201d', '"')
        .replace('\u2018', "'")
        .replace('\u2019', "'")
        .replace('\u00a0', ' ')
    )


def _parse_json_summary(summary):
    if isinstance(summary, dict):
        return summary
    text = _normalize_json_like_text(_strip_code_fence(summary))
    try:
        return json.loads(text)
    except Exception:
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except Exception:
                return None
    return None


def _find_value(data, aliases):
    if not isinstance(data, dict):
        return None
    for alias in aliases:
        if alias in data:
            return data[alias]
    normalized = {
        str(key).lower().replace('-', '_').replace(' ', '_'): value
        for key, value in data.items()
    }
    for alias in aliases:
        if alias in normalized:
            return normalized[alias]
    return None


def _coerce_list(value):
    if not value:
        return []
    if isinstance(value, list):
        return [_clean_text(item) for item in value if _clean_text(item)]
    if isinstance(value, dict):
        return [_clean_text(item) for item in value.values() if _clean_text(item)]
    return [
        re.sub(r'^[-*•\d.)\s]+', '', line).strip()
        for line in str(value).splitlines()
        if re.sub(r'^[-*•\d.)\s]+', '', line).strip()
    ]


def parse_summary_sections(summary):
    """Return normalized summary sections without exposing raw JSON."""
    data = _parse_json_summary(summary)
    if not data:
        return {
            'overview': _clean_text(summary),
            'key_points': [],
            'action_items': [],
            'decisions': [],
            'next_steps': [],
            'risks': [],
        }

    parsed = {
        'overview': _clean_text(_find_value(data, SECTION_ALIASES['overview'])),
        'key_points': _coerce_list(_find_value(data, SECTION_ALIASES['key_points'])),
        'action_items': _coerce_list(_find_value(data, SECTION_ALIASES['action_items'])),
        'decisions': _coerce_list(_find_value(data, SECTION_ALIASES['decisions'])),
        'next_steps': _coerce_list(_find_value(data, SECTION_ALIASES['next_steps'])),
        'risks': _coerce_list(_find_value(data, SECTION_ALIASES['risks'])),
    }
    if not any([
        parsed['overview'],
        parsed['key_points'],
        parsed['action_items'],
        parsed['decisions'],
        parsed['next_steps'],
        parsed['risks'],
    ]):
        parsed['overview'] = _clean_text(data)
    return parsed


def is_structured_summary(summary):
    """Return True when summary text is parseable structured JSON."""
    return _parse_json_summary(summary) is not None


def render_summary_markdown(summary, title=None):
    """Render a structured JSON summary as readable Markdown.

    Plain Markdown summaries are returned unchanged so existing PMM summaries
    keep their original formatting.
    """
    if not is_structured_summary(summary):
        return summary or ''

    sections = parse_summary_sections(summary)
    lines = ['# Meeting Summary']
    if title:
        lines.extend(['', f'**{title}**'])

    _append_markdown_text(lines, 'Summary', sections['overview'])
    _append_markdown_list(lines, 'Key Discussion Points', sections['key_points'])
    _append_markdown_list(lines, 'Action Items', sections['action_items'], marker='- [ ]')
    _append_markdown_list(lines, 'Decisions', sections['decisions'])
    _append_markdown_list(lines, 'Next Steps', sections['next_steps'])
    _append_markdown_list(lines, 'Risks', sections['risks'])
    return '\n'.join(lines).strip()


def _append_markdown_text(lines, title, text):
    if text:
        lines.extend(['', f'## {title}', '', text])


def _append_markdown_list(lines, title, items, marker='-'):
    if items:
        lines.extend(['', f'## {title}', ''])
        lines.extend(f'{marker} {item}' for item in items)


def build_summary_export_data(recording, current_user=None):
    duration = None
    try:
        duration = recording.get_audio_duration(allow_probe_fallback=False)
    except Exception:
        duration = None

    meeting_date = recording.meeting_date or recording.created_at
    sections = parse_summary_sections(recording.summary)
    return {
        'title': recording.title or 'Untitled Recording',
        'meeting_date': meeting_date.strftime('%Y-%m-%d') if meeting_date else '',
        'duration': _format_duration(duration) if duration else '',
        'participants': recording.participants or '',
        'sections': sections,
    }


def _format_duration(seconds):
    try:
        seconds = int(float(seconds))
    except Exception:
        return ''
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f'{hours}h {minutes}m {seconds}s'
    if minutes:
        return f'{minutes}m {seconds}s'
    return f'{seconds}s'


def render_summary_txt(data):
    lines = ['Meeting Summary', data['title'], '']
    info = []
    if data.get('meeting_date'):
        info.append(f"Date: {data['meeting_date']}")
    if data.get('duration'):
        info.append(f"Duration: {data['duration']}")
    if data.get('participants'):
        info.append(f"Participants: {data['participants']}")
    if info:
        lines.extend(['Meeting Information', *info, ''])

    sections = data['sections']
    _append_text(lines, 'Summary', sections['overview'])
    _append_list(lines, 'Key Discussion Points', sections['key_points'])
    _append_list(lines, 'Action Items', sections['action_items'], marker='[ ]')
    _append_list(lines, 'Decisions', sections['decisions'])
    _append_list(lines, 'Next Steps', sections['next_steps'])
    _append_list(lines, 'Risks', sections['risks'])
    lines.extend(['', 'Generated by AI Meeting Assistant'])
    return re.sub(r'\n{3,}', '\n\n', '\n'.join(lines)).strip() + '\n'


def _append_text(lines, title, text):
    if text:
        lines.extend([title, text, ''])


def _append_list(lines, title, items, marker='-'):
    if items:
        lines.extend([title, *[f'{marker} {item}' for item in items], ''])


def render_summary_docx(data):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading('Meeting Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(data['title'])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _docx_info(doc, data)
    sections = data['sections']
    _docx_text(doc, 'Summary', sections['overview'])
    _docx_list(doc, 'Key Discussion Points', sections['key_points'])
    _docx_list(doc, 'Action Items', sections['action_items'], style='List Bullet')
    _docx_list(doc, 'Decisions', sections['decisions'])
    _docx_list(doc, 'Next Steps', sections['next_steps'])
    _docx_list(doc, 'Risks', sections['risks'])
    doc.add_paragraph('Generated by AI Meeting Assistant')

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _docx_info(doc, data):
    rows = [
        ('Date', data.get('meeting_date')),
        ('Duration', data.get('duration')),
        ('Participants', data.get('participants')),
    ]
    rows = [(label, value) for label, value in rows if value]
    if not rows:
        return
    doc.add_heading('Meeting Information', level=1)
    for label, value in rows:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'{label}: ').bold = True
        paragraph.add_run(value)


def _docx_text(doc, title, text):
    if text:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)


def _docx_list(doc, title, items, style='List Bullet'):
    if not items:
        return
    doc.add_heading(title, level=1)
    for item in items:
        doc.add_paragraph(item, style=style)


def render_summary_pdf(data):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    stream = BytesIO()
    font_name, bold_font_name = _register_pdf_fonts()
    doc = SimpleDocTemplate(
        stream,
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = font_name
        style.splitLongWords = 1
        style.wordWrap = 'CJK'
    styles['Title'].fontName = bold_font_name
    styles['Heading2'].fontName = bold_font_name
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading2'],
        fontName=bold_font_name,
        spaceBefore=14,
        spaceAfter=8,
        splitLongWords=1,
        wordWrap='CJK',
    ))
    styles.add(ParagraphStyle(
        name='BodyClean',
        parent=styles['BodyText'],
        fontName=font_name,
        leading=15,
        spaceAfter=6,
        splitLongWords=1,
        wordWrap='CJK',
    ))
    styles.add(ParagraphStyle(
        name='BulletClean',
        parent=styles['BodyClean'],
        leftIndent=16,
        firstLineIndent=-10,
    ))
    styles.add(ParagraphStyle(
        name='TableCell',
        parent=styles['BodyClean'],
        fontSize=8.5,
        leading=11,
        spaceAfter=0,
    ))
    styles.add(ParagraphStyle(
        name='TableHeader',
        parent=styles['TableCell'],
        fontName=bold_font_name,
        textColor=colors.HexColor('#334155'),
    ))

    story = [
        Paragraph('Meeting Summary', styles['Title']),
        Paragraph(_escape_pdf(data['title']), styles['BodyClean']),
        Spacer(1, 12),
    ]

    info_rows = [[label, value] for label, value in [
        ('Date', data.get('meeting_date')),
        ('Duration', data.get('duration')),
        ('Participants', data.get('participants')),
    ] if value]
    if info_rows:
        story.append(Paragraph('Meeting Information', styles['SectionTitle']))
        table = Table(
            [[
                Paragraph(_escape_pdf(label), styles['TableHeader']),
                Paragraph(_escape_pdf(value), styles['TableCell']),
            ] for label, value in info_rows],
            colWidths=[1.45 * inch, 4.7 * inch],
            splitByRow=1,
            splitInRow=1,
        )
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#cbd5e1')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.extend([table, Spacer(1, 8)])

    sections = data['sections']
    _pdf_text(story, styles, 'Summary', sections['overview'])
    _pdf_list(story, styles, 'Key Discussion Points', sections['key_points'])
    _pdf_list(story, styles, 'Action Items', sections['action_items'], marker='[ ]')
    _pdf_list(story, styles, 'Decisions', sections['decisions'])
    _pdf_list(story, styles, 'Next Steps', sections['next_steps'])
    _pdf_list(story, styles, 'Risks', sections['risks'])
    story.extend([Spacer(1, 18), Paragraph('Generated by AI Meeting Assistant', styles['BodyClean'])])

    try:
        doc.build(story)
    except Exception:
        logger.exception('ReportLab failed while rendering summary PDF')
        raise
    stream.seek(0)
    return stream


def _pdf_text(story, styles, title, text):
    from reportlab.platypus import Paragraph

    if text:
        story.append(Paragraph(title, styles['SectionTitle']))
        _append_pdf_markdown(story, styles, text)


def _pdf_list(story, styles, title, items, marker='•'):
    from reportlab.platypus import Paragraph

    if not items:
        return
    story.append(Paragraph(title, styles['SectionTitle']))
    for item in items:
        story.append(Paragraph(f'{_escape_pdf(marker)} {_escape_pdf(item)}', styles['BulletClean']))


def _append_pdf_markdown(story, styles, text):
    """Append a conservative Markdown subset as ReportLab flowables."""
    from reportlab.platypus import Paragraph

    lines = str(text or '').replace('\r\n', '\n').replace('\r', '\n').split('\n')
    paragraph = []
    i = 0

    def flush_paragraph():
        if not paragraph:
            return
        paragraph_text = ' '.join(part.strip() for part in paragraph if part.strip())
        if paragraph_text:
            story.append(Paragraph(_escape_pdf(paragraph_text), styles['BodyClean']))
        paragraph.clear()

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            flush_paragraph()
            i += 1
            continue

        table = _collect_markdown_table(lines, i)
        if table:
            flush_paragraph()
            _append_pdf_table(story, styles, table)
            i += len(table) + 1
            continue

        heading = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading:
            flush_paragraph()
            story.append(Paragraph(_escape_pdf(heading.group(2)), styles['SectionTitle']))
            i += 1
            continue

        checkbox = re.match(r'^-\s+\[( |x|X)\]\s+(.+)$', line)
        if checkbox:
            flush_paragraph()
            marker = '[x]' if checkbox.group(1).lower() == 'x' else '[ ]'
            story.append(Paragraph(f'{marker} {_escape_pdf(checkbox.group(2))}', styles['BulletClean']))
            i += 1
            continue

        bullet = re.match(r'^([-*+]|(?:\d+\.))\s+(.+)$', line)
        if bullet:
            flush_paragraph()
            marker = bullet.group(1) if bullet.group(1).endswith('.') else '•'
            story.append(Paragraph(f'{_escape_pdf(marker)} {_escape_pdf(bullet.group(2))}', styles['BulletClean']))
            i += 1
            continue

        paragraph.append(line)
        i += 1

    flush_paragraph()


def _collect_markdown_table(lines, start):
    if start + 1 >= len(lines):
        return None
    header = lines[start].strip()
    separator = lines[start + 1].strip()
    separator_pattern = r'^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$'
    if not _is_markdown_table_row(header) or not re.match(separator_pattern, separator):
        return None

    rows = [_split_markdown_table_row(header)]
    index = start + 2
    while index < len(lines) and _is_markdown_table_row(lines[index].strip()):
        rows.append(_split_markdown_table_row(lines[index].strip()))
        index += 1
    return rows


def _is_markdown_table_row(line):
    return bool(line and '|' in line and len(_split_markdown_table_row(line)) >= 2)


def _split_markdown_table_row(line):
    cleaned = line.strip().strip('|')
    return [cell.strip() for cell in cleaned.split('|')]


def _append_pdf_table(story, styles, rows):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    if not rows:
        return
    max_columns = max(len(row) for row in rows)
    padded_rows = [row + [''] * (max_columns - len(row)) for row in rows]
    column_width = (6.1 * inch) / max_columns
    table_data = []
    for row_index, row in enumerate(padded_rows):
        style = styles['TableHeader'] if row_index == 0 else styles['TableCell']
        table_data.append([Paragraph(_escape_pdf(cell), style) for cell in row])

    table = Table(
        table_data,
        colWidths=[column_width] * max_columns,
        repeatRows=1,
        splitByRow=1,
        splitInRow=1,
    )
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
        ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#cbd5e1')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.extend([table, Spacer(1, 8)])


def _register_pdf_fonts():
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        regular = _find_font_file([
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/local/share/fonts/DejaVuSans.ttf',
            'C:\\Windows\\Fonts\\DejaVuSans.ttf',
            'C:\\Windows\\Fonts\\arial.ttf',
        ])
        bold = _find_font_file([
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
            '/usr/local/share/fonts/DejaVuSans-Bold.ttf',
            'C:\\Windows\\Fonts\\DejaVuSans-Bold.ttf',
            'C:\\Windows\\Fonts\\arialbd.ttf',
        ])
        if regular:
            pdfmetrics.registerFont(TTFont('PMMUnicode', regular))
            if bold:
                pdfmetrics.registerFont(TTFont('PMMUnicode-Bold', bold))
            else:
                pdfmetrics.registerFont(TTFont('PMMUnicode-Bold', regular))
            return 'PMMUnicode', 'PMMUnicode-Bold'
    except Exception:
        logger.warning('Unable to register Unicode PDF fonts; falling back to Helvetica', exc_info=True)
    return 'Helvetica', 'Helvetica-Bold'


def _find_font_file(candidates):
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return None


def _escape_pdf(text):
    return html.escape(str(text or ''), quote=False).replace('\n', '<br/>')
